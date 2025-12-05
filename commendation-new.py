import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn  # 用于设置中文字体
import io

st.title("通报表扬生成工具")

# 上传文件
uploaded_file = st.file_uploader("上传Excel文件", type=['xlsx', 'xls'])

if uploaded_file:
    try:
        # 智能读取Excel - 第一行找不到"姓名"就去第二行读
        # 先正常读取
        df_first = pd.read_excel(uploaded_file)
        
        # 检查第一行是否有"姓名"列
        has_name_in_first = any("姓名" in str(col) or "名字" in str(col) for col in df_first.columns)
        
        if has_name_in_first:
            df = df_first
            st.info("✅ 第一行找到姓名列")
        else:
            # 第一行没有"姓名"列，重新读取，跳过第一行
            uploaded_file.seek(0)  # 重置文件指针
            df = pd.read_excel(uploaded_file, header=1)
            st.info("✅ 第二行找到姓名列")
        
        # 查找姓名列
        name_column = None
        for col in df.columns:
            col_str = str(col)
            if "姓名" in col_str or "名字" in col_str:
                name_column = col
                break
        
        # 如果还是没找到，让用户选择
        if not name_column:
            st.warning("请手动选择姓名列")
            name_column = st.selectbox("选择姓名列：", df.columns)
        else:
            st.success(f"自动识别到姓名列：'{name_column}'")
        
        # 提取姓名
        if name_column:
            names = df[name_column].dropna().astype(str).str.strip().tolist()
            
            # 过滤无效数据
            names = [name for name in names if name and name != 'nan' and name != 'None']
            
            if not names:
                st.error("没有找到有效的姓名数据")
                st.stop()
            
            st.success(f"✅ 提取到 {len(names)} 个姓名")
            
            # 显示前几个姓名
            with st.expander("查看姓名预览"):
                cols = st.columns(3)
                for i, name in enumerate(names[:15]):
                    with cols[i % 3]:
                        st.write(f"{i+1}. {name}")
                if len(names) > 15:
                    st.write(f"... 等共 {len(names)} 个姓名")
            
            # 文档设置
            st.subheader("文档设置")
            
            col1, col2 = st.columns(2)
            with col1:
                # 每行姓名数量，最大10个
                per_row = st.selectbox("每行姓名数", [2, 3, 4, 5, 6, 7, 8, 9, 10], index=2)
            with col2:
                font_size = st.selectbox("姓名字体大小", [12, 14, 16], index=1)
            
            # 活动信息
            st.subheader("活动信息")
            col1, col2, col3 = st.columns(3)
            with col1:
                year = st.text_input("年份", "二〇二四")
            with col2:
                month = st.text_input("月份", "十")
            with col3:
                day = st.text_input("日期", "二十五")
            
            activity = st.text_input("活动名称", "校园文化节")
            
            # 生成文档
            if st.button("生成通报表扬"):
                with st.spinner("生成中..."):
                    # 创建文档
                    doc = Document()
                    
                    # 设置文档默认字体为宋体
                    style = doc.styles['Normal']
                    style.font.name = '宋体'
                    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    style.font.size = Pt(12)
                    
                    # 标题 - 黑体大字居中
                    title = doc.add_paragraph()
                    title_run = title.add_run("通报表扬")
                    # 标题用黑体
                    title_run.font.name = '黑体'
                    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    title_run.font.size = Pt(28)
                    title_run.bold = True
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    doc.add_paragraph()  # 空行
                    
                    # 正文 - 第一行：左对齐
                    line1 = doc.add_paragraph()
                    line1_run = line1.add_run("各学院团委及学生会：")
                    line1_run.font.name = '宋体'
                    line1_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    line1_run.font.size = Pt(14)
                    line1_run.bold = True
                    
                    # 正文 - 第二行：首行缩进
                    line2 = doc.add_paragraph()
                    line2.paragraph_format.first_line_indent = Inches(0.5)
                    line2_text = f"兹有 {year}年 {month}月 {day}日温州理工学院 {activity}活动，在以下同学的共同努力下，本次活动取得了圆满成功，经研究决定，特给予以下同学通报表扬一次："
                    line2_run = line2.add_run(line2_text)
                    line2_run.font.name = '宋体'
                    line2_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    line2_run.font.size = Pt(14)
                    
                    # 正文 - 第三行：左对齐
                    line3 = doc.add_paragraph()
                    line3_run = line3.add_run("具体名单如下：")
                    line3_run.font.name = '宋体'
                    line3_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    line3_run.font.size = Pt(14)
                    line3_run.bold = True
                    
                    doc.add_paragraph()  # 空行
                    
                    # 创建姓名表格
                    total = len(names)
                    rows = (total + per_row - 1) // per_row
                    table = doc.add_table(rows=rows, cols=per_row)
                    
                    # 填充姓名并居中
                    idx = 0
                    for row in table.rows:
                        for cell in row.cells:
                            if idx < total:
                                cell.text = names[idx]
                                paragraph = cell.paragraphs[0]
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                
                                # 设置单元格内文字的字体
                                for run in paragraph.runs:
                                    run.font.name = '宋体'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                    run.font.size = Pt(font_size)
                                
                                idx += 1
                    
                    doc.add_paragraph()  # 空行
                    
                    # 落款 - 右对齐，字体大一点
                    footer = doc.add_paragraph()
                    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    # 第一行落款
                    footer_run1 = footer.add_run("共青团温州理工学院委员会")
                    footer_run1.font.name = '宋体'
                    footer_run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    footer_run1.font.size = Pt(16)
                    footer_run1.bold = True
                    footer.add_run("\n")
                    
                    # 第二行落款
                    footer_run2 = footer.add_run(f"{year}年{month}月{day}日")
                    footer_run2.font.name = '宋体'
                    footer_run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    footer_run2.font.size = Pt(15)
                    
                    # 保存文件
                    bio = io.BytesIO()
                    doc.save(bio)
                    bio.seek(0)
                    
                    st.success("✅ 通报表扬文档生成成功！")
                    
                    # 预览文档格式
                    with st.expander("预览文档格式"):
                        st.markdown("""
                        <div style="font-family: 'SimSun', serif;">
                        <div style="text-align: center; font-size: 28px; font-weight: bold; font-family: 'SimHei';">通报表扬</div>
                        <br>
                        <div style="font-size: 14px;">
                        <div><b>各学院团委及学生会：</b></div>
                        <div style="text-indent: 2em;">
                            兹有 二〇二四年 十月 二十五日温州理工学院 校园文化节活动，在以下同学的共同努力下，本次活动取得了圆满成功，经研究决定，特给予以下同学通报表扬一次：
                        </div>
                        <div><b>具体名单如下：</b></div>
                        </div>
                        <br>
                        <div style="text-align: center; font-size: 14px;">
                        张三　　李四　　王五　　赵六<br>
                        钱七　　孙八　　周九　　吴十
                        </div>
                        <br>
                        <div style="text-align: right; font-size: 16px;">
                        <b>共青团温州理工学院委员会</b><br>
                        二〇二四年十月二十五日
                        </div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    # 下载按钮
                    st.download_button(
                        "📥 下载Word文档",
                        bio,
                        f"通报表扬_{activity}.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
    except Exception as e:
        st.error(f"处理文件时出错：{str(e)}")
        st.write("请确保上传的是正确的Excel文件")

else:
    st.info("请上传Excel文件开始使用")
    st.markdown("---")